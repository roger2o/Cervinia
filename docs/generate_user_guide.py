#!/usr/bin/env python3
"""Generate the Ski Route Planner user guide as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_screenshot_placeholder(doc, caption: str):
    """Add a bordered placeholder box indicating where to insert a screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[Insert screenshot: {caption}]\n\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    # Border the cell
    from docx.oxml.ns import qn
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = tc_pr.makeelement(qn('w:tcBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        el = tc_borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'CCCCCC',
        })
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def build_guide(include_placeholders: bool = True):
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # =========================================================
    # TITLE PAGE
    # =========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Ski Route Planner')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('User Guide')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0  |  February 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # =========================================================
    # TABLE OF CONTENTS (manual)
    # =========================================================
    doc.add_heading('Contents', level=1)
    toc_items = [
        '1. Getting Started',
        '2. Planning a Route',
        '3. The Route Panel',
        '4. The Map',
        '5. Switching Ski Areas',
        '6. GPS & Live Location',
        '7. Daily Activity Tracking',
        '8. Run History & Season Summary',
        '9. Weather Forecast',
        '10. Piste & Lift Status',
        '11. Sharing a Route',
        '12. Offline Use',
        '13. Tips & Troubleshooting',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # =========================================================
    # 1. GETTING STARTED
    # =========================================================
    doc.add_heading('1. Getting Started', level=1)
    doc.add_paragraph(
        'The Ski Route Planner is a mobile-friendly web app for planning ski routes across '
        'multiple ski areas. It calculates step-by-step directions from any station to any '
        'other station, taking into account lift connections, piste difficulty, and current '
        'lift/piste closures.'
    )
    doc.add_heading('Supported Ski Areas', level=2)
    doc.add_paragraph(
        'The app currently supports two ski areas:'
    )
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light List Accent 1'
    for i, (h) in enumerate(['Ski Area', 'Sub-Areas', 'Coverage']):
        table.rows[0].cells[i].text = h
    for i, (name, subs, coverage) in enumerate([
        ('Matterhorn Ski Paradise', 'Cervinia, Valtournenche, Zermatt', '233 stations, 158 lifts, 240 pistes'),
        ('Pontedilegno-Tonale', 'Ponte di Legno, Passo Tonale, Presena', '57 stations, 62 lifts, 65 pistes'),
    ], start=1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = subs
        table.rows[i].cells[2].text = coverage

    doc.add_paragraph()
    doc.add_heading('Opening the App', level=2)
    doc.add_paragraph(
        'Open the app in your mobile or desktop browser. On first load you will see the '
        'Matterhorn Ski Paradise map with station markers. The app can be installed as a '
        'home-screen shortcut on iOS and Android for quick access.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'App on first load showing the map with station markers')

    # =========================================================
    # 2. PLANNING A ROUTE
    # =========================================================
    doc.add_page_break()
    doc.add_heading('2. Planning a Route', level=1)

    doc.add_heading('Adding Stops', level=2)
    doc.add_paragraph(
        'There are two ways to add stops to your route:'
    )
    doc.add_paragraph('Tap a station marker on the map. The station is added as the next stop.', style='List Bullet')
    doc.add_paragraph(
        'Use the "Add stop" search box at the top. Type a station name to filter, then tap to select. '
        'Stations are grouped by sub-area for easy browsing.',
        style='List Bullet',
    )
    doc.add_paragraph(
        'You need at least two stops (a start and a destination) before a route is calculated. '
        'You can add as many intermediate stops as you like to create a multi-leg route.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Waypoint list with 3 stops and the search dropdown open')

    doc.add_heading('Reordering & Removing Stops', level=2)
    doc.add_paragraph(
        'Each stop in the list has controls to reorder or remove it:'
    )
    doc.add_paragraph('Tap the up/down arrows to move a stop earlier or later in the route.', style='List Bullet')
    doc.add_paragraph('Tap the X button to remove a stop.', style='List Bullet')

    doc.add_heading('Setting Difficulty Preference', level=2)
    doc.add_paragraph(
        'Below the waypoint list, choose your difficulty preference. This controls which pistes '
        'the router will use:'
    )
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light List Accent 1'
    table.rows[0].cells[0].text = 'Setting'
    table.rows[0].cells[1].text = 'Behaviour'
    for i, (label, desc) in enumerate([
        ('Blue', 'Only blue (easy) pistes. Safest option for beginners.'),
        ('Blue/Red', 'Prefers blue pistes but will use red if no blue route exists.'),
        ('Red', 'Uses blue and red pistes freely. The default setting.'),
        ('Red/Black', 'Prefers red or easier, but will use black if needed.'),
        ('Black', 'Uses all pistes including black (expert) runs.'),
    ], start=1):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()
    doc.add_paragraph(
        'If no route can be found at your chosen difficulty, a red error message will appear '
        'telling you which leg failed and suggesting you try a higher difficulty.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Difficulty selector with "Red" selected and a calculated route visible')

    # =========================================================
    # 3. THE ROUTE PANEL
    # =========================================================
    doc.add_page_break()
    doc.add_heading('3. The Route Panel', level=1)
    doc.add_paragraph(
        'When a route is calculated, a panel slides up from the bottom of the screen showing '
        'the step-by-step directions.'
    )

    doc.add_heading('Route Summary', level=2)
    doc.add_paragraph('The panel header shows key metrics for your route:')
    doc.add_paragraph('Total distance (km) — including lift sections', style='List Bullet')
    doc.add_paragraph('Skiing distance (km) — piste sections only', style='List Bullet')
    doc.add_paragraph('Vertical drop (m) — total descent', style='List Bullet')
    doc.add_paragraph('Duration (min) — estimated total time', style='List Bullet')
    doc.add_paragraph('Maximum difficulty — shown as a coloured dot', style='List Bullet')

    doc.add_heading('Step-by-Step Directions', level=2)
    doc.add_paragraph(
        'Each step shows an instruction such as "Take gondola Progetto Cabinovia to Station X" '
        'or "Ski down Alpino (red) to Station Y", along with the distance and duration for '
        'that segment.'
    )
    doc.add_paragraph(
        'Tap any step to highlight it on the map. The map will automatically zoom to show '
        'that segment. Tap again to deselect.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Route panel showing step-by-step directions with one step highlighted')

    doc.add_heading('Dragging the Panel', level=2)
    doc.add_paragraph(
        'The panel has three positions: collapsed (just the header), half-expanded (default), '
        'and fully expanded. Drag the handle bar at the top of the panel to resize it, or tap '
        'the header when collapsed to expand it.'
    )

    doc.add_heading('Action Buttons', level=2)
    doc.add_paragraph('Share — send the route via WhatsApp (see Section 11).', style='List Bullet')
    doc.add_paragraph('Finished — log the route to your history (see Section 8).', style='List Bullet')
    doc.add_paragraph('Clear — remove the current route and start over.', style='List Bullet')

    # =========================================================
    # 4. THE MAP
    # =========================================================
    doc.add_page_break()
    doc.add_heading('4. The Map', level=1)
    doc.add_paragraph(
        'The map uses OpenTopoMap, a topographic map that shows terrain contours, which is '
        'ideal for understanding the mountain layout.'
    )
    doc.add_heading('What You See on the Map', level=2)
    doc.add_paragraph('Station markers — white circles with a blue outline. Tap to add as a stop.', style='List Bullet')
    doc.add_paragraph('Pistes — coloured lines: blue, red, or dark grey/black matching their difficulty.', style='List Bullet')
    doc.add_paragraph('Lifts — dashed grey lines.', style='List Bullet')
    doc.add_paragraph(
        'Route overlay — when a route is active, the steps are drawn as thicker coloured lines '
        'on top of the base map. Lifts appear as dashed lines, pistes as solid lines coloured by difficulty.',
        style='List Bullet',
    )
    doc.add_paragraph(
        'Numbered stop markers — blue circles numbered 1, 2, 3… showing your route stops.',
        style='List Bullet',
    )
    doc.add_paragraph(
        'Closed indicators — red "X" markers on any closed lifts or pistes included in your route.',
        style='List Bullet',
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Map showing pistes, lifts, station markers, and a highlighted route')

    doc.add_heading('Hovering Over Markers', level=2)
    doc.add_paragraph(
        'Hover over (or long-press on mobile) any station marker to see a tooltip with the '
        'station name and elevation.'
    )

    # =========================================================
    # 5. SWITCHING SKI AREAS
    # =========================================================
    doc.add_page_break()
    doc.add_heading('5. Switching Ski Areas', level=1)
    doc.add_paragraph(
        'To switch between ski areas:'
    )
    doc.add_paragraph('Tap the hamburger menu (three lines) in the top-right corner.', style='List Number')
    doc.add_paragraph('Tap "Ski Site".', style='List Number')
    doc.add_paragraph(
        'You will see a list of available ski areas. Use the search box to filter, then tap '
        'the area you want.',
        style='List Number',
    )
    doc.add_paragraph(
        'When you switch areas, the map flies to the new area, all pistes and lifts update, '
        'and your current route is cleared. If you had previously cached a different area for '
        'offline use, you will be asked to confirm before switching.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Ski Site selector showing Matterhorn and Pontedilegno-Tonale')

    # =========================================================
    # 6. GPS
    # =========================================================
    doc.add_page_break()
    doc.add_heading('6. GPS & Live Location', level=1)
    doc.add_paragraph(
        'Tap the location button in the bottom-right corner of the map to enable GPS tracking. '
        'Your position appears as a blue dot with an accuracy circle around it.'
    )
    doc.add_paragraph(
        'When GPS is first activated, the map will fly to your current position. The button '
        'turns blue while GPS is active. Tap again to stop tracking.'
    )
    doc.add_paragraph(
        'Your browser will ask for location permission the first time. If permission is denied, '
        'an error message will appear near the button.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Map with GPS position shown as blue dot with accuracy circle')

    # =========================================================
    # 7. DAILY ACTIVITY
    # =========================================================
    doc.add_page_break()
    doc.add_heading('7. Daily Activity Tracking', level=1)
    doc.add_paragraph(
        'Track your skiing day with the built-in activity recorder. Access it from the '
        'hamburger menu under "Daily Activity".'
    )

    doc.add_heading('Recording', level=2)
    doc.add_paragraph(
        'Tap "Start Recording" to begin tracking your movement via GPS. The app will '
        'automatically enable GPS if it is not already active. While recording:'
    )
    doc.add_paragraph('A red pulsing dot appears next to "Daily Activity" in the menu.', style='List Bullet')
    doc.add_paragraph('Your total distance and maximum speed update in real time.', style='List Bullet')
    doc.add_paragraph('Data is saved automatically every 10 GPS points and when you stop recording.', style='List Bullet')
    doc.add_paragraph(
        'Tap "Stop Recording" when you are done. Your data is saved for the day and persists '
        'if you close and reopen the app.'
    )

    doc.add_heading('Viewing Your Track on the Map', level=2)
    doc.add_paragraph(
        'Check "Show track on map" to see your recorded track overlaid on the map. The track '
        'is colour-coded:'
    )
    doc.add_paragraph('Blue solid lines — skiing/downhill segments.', style='List Bullet')
    doc.add_paragraph('Grey dashed lines — lift segments (detected by speed and altitude).', style='List Bullet')

    doc.add_heading('Replay', level=2)
    doc.add_paragraph(
        'Tap "Play Replay" to watch an animated playback of your day. An orange marker moves '
        'along your track. Choose from 4 playback speeds: 1x, 2x, 5x, or 10x.'
    )

    doc.add_heading('Reset', level=2)
    doc.add_paragraph(
        'Tap "Reset" to clear the day\'s activity data. You will be asked to confirm before '
        'the data is deleted.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Daily Activity panel showing distance, speed, and track controls')

    # =========================================================
    # 8. HISTORY & SEASON
    # =========================================================
    doc.add_page_break()
    doc.add_heading('8. Run History & Season Summary', level=1)

    doc.add_heading('Logging a Run', level=2)
    doc.add_paragraph(
        'After completing a route, tap "Finished" in the route panel to log it to your history. '
        'The run is saved with the date, route details, distance, vertical drop, and duration.'
    )

    doc.add_heading('Viewing History', level=2)
    doc.add_paragraph(
        'Open the hamburger menu and tap "History" to see all your logged runs, newest first. '
        'Each entry shows the start and end stations, distance, vertical drop, and time. '
        'Tap the X to delete an individual entry.'
    )

    doc.add_heading('Season Summary', level=2)
    doc.add_paragraph(
        'Open the hamburger menu and tap "Season Summary" to see aggregate statistics for '
        'all your logged runs:'
    )
    doc.add_paragraph('Total number of runs', style='List Bullet')
    doc.add_paragraph('Total km skied', style='List Bullet')
    doc.add_paragraph('Total vertical drop (m)', style='List Bullet')
    doc.add_paragraph('Total time on the mountain', style='List Bullet')
    doc.add_paragraph('Breakdown by difficulty (blue / red / black runs)', style='List Bullet')
    doc.add_paragraph('Your top 3 most-repeated favourite routes', style='List Bullet')
    if include_placeholders: add_screenshot_placeholder(doc, 'Season Summary showing stats grid and difficulty breakdown')

    # =========================================================
    # 9. WEATHER
    # =========================================================
    doc.add_page_break()
    doc.add_heading('9. Weather Forecast', level=1)
    doc.add_paragraph(
        'Open the hamburger menu and tap "Weather" to see the current conditions and a '
        '7-day forecast for the selected ski area.'
    )
    doc.add_heading('Current Conditions', level=2)
    doc.add_paragraph(
        'Shows the current temperature, weather condition (with icon), wind speed and gusts, '
        'and the elevation of the weather station.'
    )
    doc.add_heading('7-Day Forecast', level=2)
    doc.add_paragraph(
        'Each day shows the high/low temperature, condition, precipitation or snowfall amount, '
        'and wind speed. Days are labelled "Today", "Tomorrow", or by weekday.'
    )
    doc.add_paragraph(
        'Tap "Refresh" to fetch the latest forecast. Weather data comes from the Open-Meteo '
        'service and is centred on the current ski area.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'Weather panel showing current conditions and 7-day forecast')

    # =========================================================
    # 10. P&L STATUS
    # =========================================================
    doc.add_page_break()
    doc.add_heading('10. Piste & Lift Status', level=1)
    doc.add_paragraph(
        'Open the hamburger menu and tap "P&L Status" to see which lifts and pistes are '
        'currently open or closed.'
    )
    doc.add_paragraph(
        'Each lift and piste is shown with a green "open" or red "closed" label. The header '
        'shows the count (e.g. "12/15 lifts, 30/35 pistes"). If everything is open, the '
        'panel shows "All pistes and lifts are open."'
    )
    doc.add_paragraph(
        'Closed lifts and pistes are automatically avoided by the route planner. If a closed '
        'segment must be used (no alternative exists), it will be shown with a red X marker '
        'on the map and a warning in the route panel.'
    )
    doc.add_paragraph(
        'Tap "Refresh" to fetch the latest status data.'
    )
    if include_placeholders: add_screenshot_placeholder(doc, 'P&L Status panel showing open and closed lifts')

    # =========================================================
    # 11. SHARING
    # =========================================================
    doc.add_page_break()
    doc.add_heading('11. Sharing a Route', level=1)
    doc.add_paragraph(
        'With an active route, tap the "Share" button in the route panel. This opens WhatsApp '
        'with a pre-composed message containing:'
    )
    doc.add_paragraph('Start and end station names', style='List Bullet')
    doc.add_paragraph('Number of stops (if multi-stop)', style='List Bullet')
    doc.add_paragraph('Total distance, skiing distance, vertical drop, and duration', style='List Bullet')
    doc.add_paragraph('A clickable link that recreates the exact route when opened', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph(
        'When the recipient opens the shared link, the app will automatically load the correct '
        'ski area, set the difficulty, and calculate the same route.'
    )
    doc.add_paragraph(
        'Note: Sharing requires an internet connection. If you are offline, nothing will happen.'
    )

    # =========================================================
    # 12. OFFLINE USE
    # =========================================================
    doc.add_page_break()
    doc.add_heading('12. Offline Use', level=1)
    doc.add_paragraph(
        'The app is designed to work on the mountain where signal can be patchy. Once loaded, '
        'it works fully offline.'
    )

    doc.add_heading('Automatic Caching', level=2)
    doc.add_paragraph(
        'When you first load a ski area, the route data (stations, pistes, lifts) is cached '
        'locally. On subsequent visits, the app loads instantly from the cache even without '
        'internet.'
    )

    doc.add_heading('Downloading Map Tiles', level=2)
    doc.add_paragraph(
        'The map tiles (the background topographic map) are only cached as you browse. To '
        'ensure full map coverage offline:'
    )
    doc.add_paragraph('Open the hamburger menu and tap "Offline Map".', style='List Number')
    doc.add_paragraph(
        'Tap "Download Map Tiles". This pre-downloads all tiles for the current ski area '
        'at the relevant zoom levels.',
        style='List Number',
    )
    doc.add_paragraph('Wait for the progress bar to complete.', style='List Number')
    doc.add_paragraph()
    doc.add_paragraph(
        'It is recommended to do this on Wi-Fi before heading to the mountain.'
    )

    doc.add_heading('Offline Indicator', level=2)
    doc.add_paragraph(
        'When you lose internet connection, a brief amber banner appears at the top: '
        '"You are offline — using cached data". This auto-dismisses after 2 seconds.'
    )

    doc.add_heading('What Works Offline', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light List Accent 1'
    table.rows[0].cells[0].text = 'Feature'
    table.rows[0].cells[1].text = 'Offline?'
    for i, (feat, status) in enumerate([
        ('Route planning & directions', 'Yes — fully functional'),
        ('Map display', 'Yes — if tiles were cached or pre-downloaded'),
        ('GPS & live location', 'Yes — GPS works without internet'),
        ('Daily activity tracking', 'Yes — stored locally'),
        ('Run history & season summary', 'Yes — stored locally'),
        ('Weather & P&L status', 'No — requires internet to refresh'),
    ], start=1):
        table.rows[i].cells[0].text = feat
        table.rows[i].cells[1].text = status

    if include_placeholders: add_screenshot_placeholder(doc, 'Offline Map download panel with progress bar')

    # =========================================================
    # 13. TIPS & TROUBLESHOOTING
    # =========================================================
    doc.add_page_break()
    doc.add_heading('13. Tips & Troubleshooting', level=1)

    doc.add_heading('"No route found"', level=2)
    doc.add_paragraph(
        'This means the router cannot find a path between your stops at the selected '
        'difficulty. The error message will tell you which leg failed. Try:'
    )
    doc.add_paragraph('Increasing the difficulty preference (e.g. from Blue to Red).', style='List Bullet')
    doc.add_paragraph('Checking P&L Status — a key lift may be closed.', style='List Bullet')
    doc.add_paragraph('Removing or reordering intermediate stops.', style='List Bullet')

    doc.add_heading('Map not showing pistes', level=2)
    doc.add_paragraph(
        'If pistes are not visible after switching areas, try zooming in or out to trigger '
        'a map refresh. Ensure you have an internet connection for the initial load.'
    )

    doc.add_heading('GPS not working', level=2)
    doc.add_paragraph(
        'Ensure you have granted location permission to your browser. On iOS, this is under '
        'Settings > Safari > Location. On Android, check the site permissions in Chrome.'
    )

    doc.add_heading('Installing as an App', level=2)
    doc.add_paragraph(
        'On iOS: Open in Safari, tap the share button, then "Add to Home Screen".'
    )
    doc.add_paragraph(
        'On Android: Open in Chrome, tap the three-dot menu, then "Add to Home Screen" or '
        '"Install App".'
    )

    doc.add_heading('Battery Saving', level=2)
    doc.add_paragraph(
        'GPS tracking and daily activity recording use battery. If you do not need live '
        'tracking, tap the location button to turn off GPS when not in use.'
    )

    # =========================================================
    # SAVE
    # =========================================================
    suffix = '' if include_placeholders else '_no_screenshots'
    output_path = Path(__file__).parent / f'UserGuide{suffix}.docx'
    doc.save(str(output_path))
    print(f'User guide saved to {output_path}')


if __name__ == '__main__':
    build_guide(include_placeholders=True)
    build_guide(include_placeholders=False)
